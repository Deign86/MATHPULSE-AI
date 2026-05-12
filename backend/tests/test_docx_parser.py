import pytest
from backend.rag.docx_parser import parse_docx_structure
from docx import Document
from io import BytesIO

def create_mock_docx():
    doc = Document()
    doc.add_heading('Lesson 1: Quadratic Equations', 0)
    doc.add_paragraph('Ang quadratic equation ay nasa anyong ax² + bx + c = 0.')
    doc.add_heading('Objectives', level=1)
    doc.add_paragraph('1. Identify quadratic terms.')
    
    f = BytesIO()
    doc.save(f)
    return f.getvalue()

def test_parse_docx_structure():
    content = create_mock_docx()
    elements = parse_docx_structure(content)
    
    assert len(elements) == 4
    assert elements[0]['is_heading'] is True
    assert 'Quadratic Equations' in elements[0]['text']
    assert 'Ang quadratic equation' in elements[1]['text']
    assert elements[2]['is_heading'] is True
    assert elements[2]['text'] == 'Objectives'
